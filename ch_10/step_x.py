from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Mm

from step_1_1 import OUT_DIR  # 이전에 작성한 모듈을 불러옵니다.
from step_3_1 import apply_font
from step_3_4 import OUT_3_4

OUT_X = OUT_DIR / f"{Path(__file__).stem}.docx"


def insert_info():
    doc = Document(str(OUT_3_4))  # 워드 문서 불러오기
    table = doc.add_table(rows=1, cols=1, style="Light Shading Accent 6")  # 표 추가
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 표 가로 정렬
    table.allow_autofit = False  # 표 너비 자동 맞춤 해제

    tr = table.rows[0]  # 첫 번째 행 반환
    td = tr.cells[0]  # 첫 번째 셀 반환
    td.width = Mm(174)  # 셀 너비 설정

    p1 = td.paragraphs[0]  # 첫 번째 단락 반환
    r1 = p1.add_run("참고사항")  # 참고사항 입력
    apply_font(r1, size_pt=10, is_bold=True)

    p2 = td.add_paragraph(style="List Bullet")  # 두 번째 단락 추가
    r2 = p2.add_run("주요 금리 현황은 ECOS 월단위 데이터를 기준으로 작성되었습니다.")
    apply_font(r2, size_pt=9, is_bold=False)

    p3 = td.add_paragraph(style="List Bullet")  # 세 번째 단락 추가
    r3 = p3.add_run("정기예금 상품의 금리는 수시로 변경될 수 있습니다.")
    apply_font(r3, size_pt=9, is_bold=False)
    doc.save(str(OUT_X))  # 워드 파일로 저장


if __name__ == "__main__":
    insert_info()  # 참고사항 입력
