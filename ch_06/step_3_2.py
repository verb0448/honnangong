import json
from pathlib import Path

from docx.document import Document as DocumentObject  # Document 객체
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 단락 정렬 방법
from docx.shared import Cm  # 크기 설정 클래스(센티미터)

from step_1_1 import OUT_DIR  # 이전에 작성한 모듈을 불러옵니다.
from step_2_2 import OUT_2_2
from step_2_3 import fetch_trends_by_filter
from step_3_1 import apply_font_style, init_docx


def add_table(doc: DocumentObject, category: str, option: str):
    fetch_trends_by_filter(category, option)  # 쇼핑 트렌드 데이터 수집
    imgs_path: list = json.loads(OUT_2_2.read_text(encoding="utf8"))  # 이미지 경로
    n_items = len(imgs_path)  # 이미지 개수
    n_cols = 5  # 열 개수
    n_rows = n_items // n_cols + (1 if n_items % n_cols > 0 else 0)  # 행 개수

    para = doc.add_paragraph(style="Heading 2")  # 표 제목
    text_filter = f"{option}의 {category} 트렌드"
    apply_font_style(para.add_run(text_filter), size_pt=15, is_bold=True)

    table = doc.add_table(rows=n_rows, cols=n_cols, style="Table Grid")  # 표 생성
    for tr in table.rows:  # 행 반복 처리
        for td in tr.cells:  # 열 반복 처리
            if len(imgs_path) > 0:  # 처리할 경로가 남아있는 경우
                img_path = imgs_path.pop(0)  # 첫 번째 위치의 경로 추출
                p_cell = td.paragraphs[0]  # 현재 셀의 기본 Paragraph 객체 선택
                p_cell.add_run().add_picture(img_path, width=Cm(3))  # 이미지 추가
                p_cell.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 단락 정렬

    doc.add_paragraph("")  # 빈 단락 추가
    p_list = doc.add_paragraph(style="List Bullet")  # 목록 서식 적용한 단락 추가
    text_notice = "보다 자세한 정보는 네이버플러스 스토어에서 확인하세요."
    apply_font_style(p_list.add_run(text_notice), size_pt=9)


if __name__ == "__main__":
    doc = init_docx()
    add_table(doc, "패션의류", "10대 여성")  # 표 추가
    doc.save(OUT_DIR / f"{Path(__file__).stem}.docx")
